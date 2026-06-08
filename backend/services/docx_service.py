from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from schemas.doc_style import StyleConfig

_DEFAULT_STYLE = StyleConfig()


def _hex_to_rgbcolor(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#").zfill(6)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_font(run, name: str = "Calibri", size: float = 11, bold: bool = False,
              italic: bool = False, color: str | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = _hex_to_rgbcolor(color)


def _apply_entry_subtitle_style(run, style: StyleConfig):
    es = style.entry_subtitle_style
    run.bold = es in ("bold", "bold_italic")
    run.italic = es in ("italic", "bold_italic")


def _add_section_border(paragraph, style: StyleConfig):
    """Apply section heading visual style based on section_heading_style."""
    from docx.oxml.ns import qn
    from lxml import etree

    heading_style = style.section_heading_style
    line_color = style.color_heading_line

    # Map legacy section_separator to new style when default
    if heading_style == "line_below" and style.section_separator != "line":
        sep_map = {
            "thick_line": "thick_line_below",
            "double_line": "double_line_below",
            "none": "plain",
        }
        heading_style = sep_map.get(style.section_separator, "line_below")

    if heading_style == "plain":
        return

    if heading_style == "underline":
        for run in paragraph.runs:
            run.font.underline = True
        return

    pPr = paragraph._element.get_or_add_pPr()

    if heading_style in ("line_below", "thick_line_below", "double_line_below"):
        val_map = {
            "line_below": ("single", "4"),
            "thick_line_below": ("thick", "6"),
            "double_line_below": ("double", "4"),
        }
        val, sz = val_map[heading_style]
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), val)
        bottom.set(qn("w:sz"), sz)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), line_color)

    elif heading_style == "boxed":
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        for side in ("top", "left", "bottom", "right"):
            el = etree.SubElement(pBdr, qn(f"w:{side}"))
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), line_color)

    elif heading_style == "bar":
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        left = etree.SubElement(pBdr, qn("w:left"))
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")   # ~2.25pt thick
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), line_color)


def _add_left_tabstop(paragraph, position_inches: float):
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = etree.SubElement(pPr, qn("w:tabs"))
    tab = etree.SubElement(tabs, qn("w:tab"))
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))


def _set_line_spacing(paragraph, spacing: float):
    if spacing != 1.0:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = spacing


def _add_tabstop_right(paragraph, position_inches: float):
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = etree.SubElement(pPr, qn("w:tabs"))
    tab = etree.SubElement(tabs, qn("w:tab"))
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))


def _bullet_char_for_style(style: StyleConfig) -> str | None:
    """Return bullet prefix string or None for 'none' list style."""
    if style.entry_list_style == "none":
        return None
    if style.entry_list_style == "dash":
        return "–"
    return style.bullet_char


def create_resume(
    user_name: str,
    location: str | None,
    email: str | None,
    phone: str | None,
    linkedin: str | None,
    summary: str | None,
    skills: list[dict] | None,
    educations: list[dict],
    tailored_experiences: list[dict],
    output_path: Path,
    style: StyleConfig | None = None,
) -> Path:
    """Generate a resume DOCX at output_path. Returns the path."""
    s = style or _DEFAULT_STYLE
    doc = Document()

    page_width_inches = 8.5
    content_width = page_width_inches - s.margin_left - s.margin_right

    for section in doc.sections:
        section.top_margin = Inches(s.margin_top)
        section.bottom_margin = Inches(s.margin_bottom)
        section.left_margin = Inches(s.margin_left)
        section.right_margin = Inches(s.margin_right)

    align = WD_ALIGN_PARAGRAPH.CENTER if s.header_layout == "centered" else WD_ALIGN_PARAGRAPH.LEFT

    # --- Headline ---
    name_para = doc.add_paragraph()
    name_para.alignment = align
    name_para.space_after = Pt(2)
    name_text = user_name.upper() if s.name_uppercase else user_name
    run = name_para.add_run(name_text)
    _set_font(run, name=s.font_name, size=s.font_size_name, bold=s.name_bold,
              italic=s.name_italic, color=s.name_color)
    if s.name_letter_spacing:
        from docx.oxml.ns import qn
        rPr = run._r.get_or_add_rPr()
        spacing_el = rPr.find(qn("w:spacing"))
        if spacing_el is None:
            from lxml import etree
            spacing_el = etree.SubElement(rPr, qn("w:spacing"))
        spacing_el.set(qn("w:val"), str(int(s.name_letter_spacing * 20)))

    contact_parts = [p for p in [location, phone, email, linkedin] if p]
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = align
        contact_para.space_after = Pt(6)
        run = contact_para.add_run(f" {s.contact_separator} ".join(contact_parts))
        _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_contact)

    def _section_heading(label: str):
        heading = doc.add_paragraph()
        heading.space_before = Pt(s.space_before_section)
        heading.space_after = Pt(s.space_after_section)
        text = label.upper() if s.section_caps else label
        run = heading.add_run(text)
        _set_font(run, name=s.font_name, size=s.font_size_section, bold=s.section_bold,
                  color=s.color_heading)
        _add_section_border(heading, s)
        return heading

    def _render_summary():
        if not summary:
            return
        _section_heading("Professional Summary")
        summary_para = doc.add_paragraph()
        summary_para.space_after = Pt(4)
        _set_line_spacing(summary_para, s.line_spacing)
        run = summary_para.add_run(summary)
        _set_font(run, name=s.font_name, size=s.font_size_body)

    def _render_skills():
        if not skills:
            return
        _section_heading("Technical Skills")
        for skill_cat in skills:
            skill_para = doc.add_paragraph()
            skill_para.space_after = Pt(1)
            _set_line_spacing(skill_para, s.line_spacing)
            cat_run = skill_para.add_run(f"{skill_cat['category']}: ")
            _set_font(cat_run, name=s.font_name, size=s.font_size_body, bold=True)
            skills_run = skill_para.add_run(", ".join(skill_cat["skills"]))
            _set_font(skills_run, name=s.font_name, size=s.font_size_body)

    def _render_experience():
        if not tailored_experiences:
            return
        _section_heading("Professional Experience")
        bullet_prefix = _bullet_char_for_style(s)
        indent = Inches(0.25) if s.entry_indent_body else Inches(0)

        for i, exp in enumerate(tailored_experiences):
            end = exp.get("end_date") or "Present"
            exp_location = exp.get("location") or ""
            company = exp.get("company") or ""
            title = exp.get("title") or ""

            company_text = f"{company}, {exp_location}" if exp_location else company
            dates_text = f"{exp.get('start_date') or ''} – {end}"

            if s.experience_layout == "combined":
                # Single line: "Title | Company" with dates right-aligned
                entry_para = doc.add_paragraph()
                entry_para.space_before = Pt(s.space_between_entries)
                entry_para.space_after = Pt(1)
                _add_tabstop_right(entry_para, content_width)
                combined_text = f"{title} | {company_text}"
                run = entry_para.add_run(combined_text)
                _set_font(run, name=s.font_name, size=s.entry_title_size, bold=True,
                          color=s.color_job_title)
                run = entry_para.add_run(f"\t{dates_text}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_dates)

            elif s.experience_layout == "title-employer":
                # Title (bold) on line 1, employer on line 2
                title_para = doc.add_paragraph()
                title_para.space_before = Pt(s.space_between_entries)
                title_para.space_after = Pt(1)
                _add_tabstop_right(title_para, content_width)
                run = title_para.add_run(title)
                _set_font(run, name=s.font_name, size=s.entry_title_size, bold=True,
                          color=s.color_job_title)
                run = title_para.add_run(f"\t{dates_text}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_dates)

                employer_para = doc.add_paragraph()
                employer_para.space_after = Pt(1)
                run = employer_para.add_run(company_text)
                _set_font(run, name=s.font_name, size=s.entry_subtitle_size,
                          color=s.color_employer)
                _apply_entry_subtitle_style(run, s)

            else:
                # employer-title (default): Company bold on line 1, title + dates on line 2
                company_para = doc.add_paragraph()
                company_para.space_before = Pt(s.space_between_entries)
                company_para.space_after = Pt(1)
                run = company_para.add_run(company_text)
                _set_font(run, name=s.font_name, size=s.entry_subtitle_size, bold=True,
                          color=s.color_employer)

                title_para = doc.add_paragraph()
                title_para.space_after = Pt(1)
                _add_tabstop_right(title_para, content_width)
                run = title_para.add_run(title)
                _set_font(run, name=s.font_name, size=s.entry_title_size, color=s.color_job_title)
                run = title_para.add_run(f"\t{dates_text}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_dates)

            for bullet in exp.get("bullets", []):
                bullet_para = doc.add_paragraph()
                bullet_para.space_after = Pt(1)
                if s.entry_indent_body:
                    bullet_para.paragraph_format.left_indent = indent
                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                    _add_left_tabstop(bullet_para, 0.25)
                _set_line_spacing(bullet_para, s.line_spacing)
                if bullet_prefix:
                    run = bullet_para.add_run(f"{bullet_prefix}\t{bullet}")
                else:
                    run = bullet_para.add_run(bullet)
                _set_font(run, name=s.font_name, size=s.font_size_body)

    def _render_education():
        if not educations:
            return
        _section_heading("Education")
        for edu in educations:
            end = edu.get("end_date") or "Present"

            edu_para = doc.add_paragraph()
            edu_para.space_before = Pt(s.space_between_entries)
            edu_para.space_after = Pt(1)
            _add_tabstop_right(edu_para, content_width)

            degree_text = f"{edu['degree']} in {edu['field']}"
            school_text = edu["school"]
            dates_text = f"{edu['start_date']} – {end}"

            if s.education_layout == "school-degree":
                # School bold on top, degree below
                run = edu_para.add_run(school_text)
                _set_font(run, name=s.font_name, size=s.entry_title_size, bold=True,
                          color=s.color_employer)
                run = edu_para.add_run(f"\t{dates_text}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_dates)

                subtitle_para = doc.add_paragraph()
                subtitle_para.space_after = Pt(1)
                run = subtitle_para.add_run(degree_text)
                _set_font(run, name=s.font_name, size=s.entry_subtitle_size,
                          color=s.color_subtitle)
                _apply_entry_subtitle_style(run, s)
            else:
                # degree-school (default): Degree bold on top, school below
                run = edu_para.add_run(f"{degree_text}, {school_text}")
                _set_font(run, name=s.font_name, size=s.entry_title_size, bold=True,
                          color=s.color_job_title)
                run = edu_para.add_run(f"\t{dates_text}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_dates)

            if edu.get("gpa"):
                gpa_para = doc.add_paragraph()
                gpa_para.space_after = Pt(1)
                run = gpa_para.add_run(f"GPA: {edu['gpa']}")
                _set_font(run, name=s.font_name, size=s.font_size_contact, color=s.color_subtitle)

    # --- Render sections in configured order ---
    _render_map = {
        "summary": _render_summary,
        "skills": _render_skills,
        "experience": _render_experience,
        "education": _render_education,
    }

    for sec in s.sections:
        if sec.visible and sec.key in _render_map:
            _render_map[sec.key]()

    doc.save(str(output_path))
    return output_path


def create_cover_letter(
    user_name: str,
    email: str | None,
    phone: str | None,
    cover_letter_text: str,
    job_title: str,
    company: str | None,
    output_path: Path,
) -> Path:
    """Generate a cover letter DOCX at output_path. Returns the path."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.space_after = Pt(2)
    run = name_para.add_run(user_name)
    _set_font(run, size=14, bold=True)

    contact_parts = [p for p in [email, phone] if p]
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.space_after = Pt(12)
        run = contact_para.add_run(" | ".join(contact_parts))
        _set_font(run, size=10)

    date_para = doc.add_paragraph()
    date_para.space_after = Pt(12)
    run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    _set_font(run, size=11)

    paragraphs = cover_letter_text.split("\n\n")
    for text in paragraphs:
        text = text.strip()
        if not text:
            continue
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        run = para.add_run(text)
        _set_font(run, size=11)

    doc.save(str(output_path))
    return output_path
